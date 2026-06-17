"""
# 1. Resume Parsing Module
Purpose: To extract clean text from standard resume formats (PDF and DOCX).
ML/NLP Concepts: Data Ingestion & Normalization. Before any NLP model can process text, 
it must be extracted cleanly, handling artifacts like headers, footers, tables, and special characters.

Architecture: A class-based parser that determines the file type and routes it to the appropriate extraction logic.
Optimization: We use `pdfplumber` instead of `PyPDF2` because it handles multi-column layouts and tables better, which are common in modern resumes.
"""

import os
import pdfplumber  # type: ignore
import docx  # type: ignore

class ResumeParser:
    def __init__(self):
        pass
        
    def parse(self, file_path: str) -> str:
        """
        Routes the file to the correct parsing method based on extension.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")
            
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext == '.docx':
            return self._parse_docx(file_path)
        elif ext == '.txt':
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")
            
    def _parse_pdf(self, file_path: str) -> str:
        """
        Extracts text from a PDF file using pdfplumber.
        Line-by-line explanation:
        1. Open the PDF file using pdfplumber's context manager (ensures it closes properly).
        2. Iterate through every page in the PDF.
        3. Extract text from the page. If text exists, add it to our list.
        4. Join all extracted text blocks with newlines.
        """
        text_content = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            return "\n".join(text_content)
        except Exception as e:
            print(f"Error parsing PDF {file_path}: {e}")
            return ""

    def _parse_docx(self, file_path: str) -> str:
        """
        Extracts text from a DOCX file using python-docx.
        """
        try:
            doc = docx.Document(file_path)
            # doc.paragraphs contains all paragraphs in the document
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            print(f"Error parsing DOCX {file_path}: {e}")
            return ""
            
    def _parse_txt(self, file_path: str) -> str:
        """
        Extracts text from a raw TXT file (useful for testing).
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error parsing TXT {file_path}: {e}")
            return ""

# Testing block to run the module directly
if __name__ == "__main__":
    parser = ResumeParser()
    print("Parser Module is ready to use.")
